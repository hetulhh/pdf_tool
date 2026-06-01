"""
PDF → Word (.docx) conversion service.
Pure logic — no Flask imports.
"""

import io
import logging

import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
# Public entry point
# ─────────────────────────────────────────────

def convert(pdf_bytes: bytes) -> bytes:
    doc = Document()
    _set_default_style(doc)
    _set_page_margins(doc)

    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        total_pages = len(pdf.pages)

        for page_num, page in enumerate(pdf.pages, start=1):

            # ── Page heading ────────────────────────────────────────
            heading = doc.add_heading(f"Page {page_num}", level=1)
            _style_heading(heading)

            # ── Tables ──────────────────────────────────────────────
            tables = page.extract_tables({
                "vertical_strategy":   "lines",
                "horizontal_strategy": "lines",
            }) or page.extract_tables({
                "vertical_strategy":   "text",
                "horizontal_strategy": "text",
            })

            inserted_tables = 0
            for raw in (tables or []):
                if not raw or len(raw) < 2:
                    continue
                _write_table(doc, raw)
                inserted_tables += 1

            # ── Text ────────────────────────────────────────────────
            text = (page.extract_text(x_tolerance=3, y_tolerance=3) or "").strip()
            if text:
                if inserted_tables:
                    sub = doc.add_heading("Text Content", level=2)
                    _style_heading(sub)
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        para = doc.add_paragraph(line)
                        para.runs[0].font.size = Pt(10)

            # ── Page break (not after last page) ────────────────────
            if page_num < total_pages:
                doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────
# Table helper
# ─────────────────────────────────────────────

def _write_table(doc: Document, raw: list):
    rows = [r for r in raw if any(c for c in r if c)]
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table    = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx in range(num_cols):
            cell_val = row[c_idx] if c_idx < len(row) else ""
            cell     = table.cell(r_idx, c_idx)
            cell.text = str(cell_val or "").strip()

            para = cell.paragraphs[0]
            if para.runs:
                run = para.runs[0]
                run.font.name = "Calibri"
                run.font.size = Pt(9)
                if r_idx == 0:
                    run.font.bold  = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    _set_cell_background(cell, "2F5496")

    doc.add_paragraph()  # spacing after table


def _set_cell_background(cell, hex_color: str):
    """Fill a table cell with a solid background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


# ─────────────────────────────────────────────
# Document style helpers
# ─────────────────────────────────────────────

def _set_default_style(doc: Document):
    style      = doc.styles["Normal"]
    font       = style.font
    font.name  = "Calibri"
    font.size  = Pt(10)


def _set_page_margins(doc: Document, margin_inches: float = 1.0):
    for section in doc.sections:
        section.top_margin    = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin   = Inches(margin_inches)
        section.right_margin  = Inches(margin_inches)


def _style_heading(heading):
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        run.font.name      = "Calibri"
